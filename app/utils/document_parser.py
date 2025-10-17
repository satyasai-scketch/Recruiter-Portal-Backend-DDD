"""
Document parsing utilities for extracting text from PDF, DOCX, and DOC files.
"""

import io
import logging
from typing import Optional, Dict, Any
from pathlib import Path

try:
    import PyPDF2
    import docx
    from docx import Document
except ImportError:
    PyPDF2 = None
    docx = None
    Document = None

logger = logging.getLogger(__name__)


class DocumentParseError(Exception):
    """Raised when document parsing fails."""
    pass


class DocumentParser:
    """Utility class for parsing various document formats."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    @classmethod
    def is_supported_format(cls, filename: str) -> bool:
        """Check if the file format is supported."""
        import logging
        logger = logging.getLogger(__name__)
        
        logger.info(f"=== is_supported_format START ===")
        logger.info(f"Filename: '{filename}'")
        logger.info(f"Filename type: {type(filename)}")
        logger.info(f"Filename repr: {repr(filename)}")
        
        if not filename:
            logger.info(f"Filename is empty/None, returning False")
            return False
        
        try:
            extension = Path(filename).suffix.lower()
            logger.info(f"File extension: '{extension}'")
            is_supported = extension in cls.SUPPORTED_EXTENSIONS
            logger.info(f"Is supported: {is_supported}")
            logger.info(f"Supported extensions: {cls.SUPPORTED_EXTENSIONS}")
            return is_supported
        except Exception as e:
            logger.error(f"Error checking file format: {str(e)}")
            logger.error(f"Error type: {type(e).__name__}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            return False
    
    @classmethod
    def validate_file_size(cls, file_size: int) -> bool:
        """Validate file size is within limits."""
        return file_size <= cls.MAX_FILE_SIZE
    
    @classmethod
    def extract_text_from_pdf(cls, file_content: bytes) -> str:
        """Extract text from PDF file content."""
        if not PyPDF2:
            raise DocumentParseError("PyPDF2 is not installed. Install with: pip install PyPDF2")
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            if not text_content:
                raise DocumentParseError("No text content found in PDF")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            raise DocumentParseError(f"Failed to parse PDF: {str(e)}")
    
    @classmethod
    def extract_text_from_docx(cls, file_content: bytes) -> str:
        """Extract text from DOCX or DOC file content."""
        if not docx:
            raise DocumentParseError("python-docx is not installed. Install with: pip install python-docx")
        
        try:
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            if not text_content:
                raise DocumentParseError("No text content found in DOCX/DOC")
            
            return "\n".join(text_content)
            
        except Exception as e:
            raise DocumentParseError(f"Failed to parse DOCX/DOC: {str(e)}")
    
    @classmethod
    def extract_text(cls, filename: str, file_content: bytes) -> Dict[str, Any]:
        """
        Extract text from document file.
        
        Returns:
            Dict containing extracted text and metadata
        """
        import logging
        logger = logging.getLogger(__name__)
        
        logger.info(f"=== DOCUMENT PARSER extract_text START ===")
        logger.info(f"Filename: '{filename}'")
        logger.info(f"Filename type: {type(filename)}")
        logger.info(f"Filename repr: {repr(filename)}")
        logger.info(f"File content size: {len(file_content)} bytes")
        
        logger.info(f"Step 1: Checking supported format...")
        if not cls.is_supported_format(filename):
            logger.error(f"Unsupported file format: {filename}")
            raise DocumentParseError(f"Unsupported file format: {filename}")
        logger.info(f"File format is supported")
        
        logger.info(f"Step 2: Validating file size...")
        if not cls.validate_file_size(len(file_content)):
            logger.error(f"File size {len(file_content)} exceeds limit {cls.MAX_FILE_SIZE}")
            raise DocumentParseError(f"File size exceeds maximum limit of {cls.MAX_FILE_SIZE} bytes")
        logger.info(f"File size is valid")
        
        logger.info(f"Step 3: Getting file extension...")
        try:
            extension = Path(filename).suffix.lower()
            logger.info(f"File extension: '{extension}'")
        except Exception as e:
            logger.error(f"Failed to get file extension: {str(e)}")
            logger.error(f"Error type: {type(e).__name__}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            raise DocumentParseError(f"Failed to get file extension from '{filename}': {str(e)}")
        
        logger.info(f"Step 4: Extracting text based on extension...")
        try:
            if extension == '.pdf':
                logger.info(f"Processing PDF file...")
                extracted_text = cls.extract_text_from_pdf(file_content)
                logger.info(f"PDF text extraction completed")
            elif extension in ['.docx', '.doc']:
                logger.info(f"Processing DOCX/DOC file...")
                extracted_text = cls.extract_text_from_docx(file_content)
                logger.info(f"DOCX/DOC text extraction completed")
            else:
                logger.error(f"Unsupported extension: {extension}")
                raise DocumentParseError(f"Unsupported file format: {extension}")
            
            logger.info(f"Step 5: Preparing result...")
            result = {
                'extracted_text': extracted_text.strip(),
                'original_filename': filename,
                'file_size': len(file_content),
                'file_extension': extension,
                'word_count': len(extracted_text.split()),
                'character_count': len(extracted_text)
            }
            logger.info(f"Result prepared successfully")
            logger.info(f"Extracted text length: {len(extracted_text)} characters")
            logger.info(f"Word count: {result['word_count']}")
            logger.info(f"=== DOCUMENT PARSER extract_text COMPLETED ===")
            return result
            
        except Exception as e:
            logger.error(f"Text extraction failed: {str(e)}")
            logger.error(f"Error type: {type(e).__name__}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            if isinstance(e, DocumentParseError):
                raise
            raise DocumentParseError(f"Failed to extract text from {filename}: {str(e)}")


def extract_job_description_text(filename: str, file_content: bytes) -> Dict[str, Any]:
    """
    Convenience function to extract text from job description documents.
    
    Args:
        filename: Name of the uploaded file
        file_content: Binary content of the file
        
    Returns:
        Dict containing extracted text and metadata
        
    Raises:
        DocumentParseError: If parsing fails
    """
    import logging
    logger = logging.getLogger(__name__)
    
    logger.info(f"=== DOCUMENT PARSER extract_job_description_text START ===")
    logger.info(f"Filename: '{filename}'")
    logger.info(f"Filename type: {type(filename)}")
    logger.info(f"Filename repr: {repr(filename)}")
    logger.info(f"File content size: {len(file_content)} bytes")
    
    try:
        result = DocumentParser.extract_text(filename, file_content)
        logger.info(f"Document parsing successful")
        logger.info(f"Result keys: {list(result.keys())}")
        logger.info(f"=== DOCUMENT PARSER extract_job_description_text COMPLETED ===")
        return result
    except Exception as e:
        logger.error(f"Document parsing failed: {str(e)}")
        logger.error(f"Error type: {type(e).__name__}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise
